import os
import sys

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx is not installed. Running with uv --with python-docx...")
    sys.exit(1)

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins on a table cell (values in dxas, 20 dxas = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    """Add a heading with consistent brand styling (Navy/Indigo colors)."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    # Custom color and font sizes
    run = h.runs[0]
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 58, 86)  # Deep Navy #1F3A56
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(99, 102, 241)  # Indigo Accent #6366F1
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(75, 85, 99)  # Slate Gray #4B5563
        run.font.bold = True
        run.font.italic = True
    return h

def set_font_style(run, name='Calibri', size=11, bold=False, italic=False, color_rgb=None):
    """Set typography details on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_paragraph_styled(doc, text="", space_after=6, line_spacing=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add body paragraph with brand standard styling."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_font_style(run)
    return p

def create_table_styled(doc, rows, cols, col_widths=None):
    """Create a beautiful branded table with padding and light borders."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Apply width if provided
    if col_widths:
        for i, col in enumerate(table.columns):
            if i < len(col_widths):
                col.width = Inches(col_widths[i])
                
    # Style borders (add thin gray grid lines)
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    return table

def populate_row(row, values, bg_color=None, is_header=False, bold=False):
    """Populate cells in a row, styling each cell."""
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)  # 6pt top/bottom, 7.5pt left/right padding
        if bg_color:
            set_cell_background(cell, bg_color)
            
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        
        run = p.add_run(str(val))
        if is_header:
            set_font_style(run, size=10, bold=True, color_rgb=RGBColor(255, 255, 255))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            set_font_style(run, size=9.5, bold=bold)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def apply_bullet(doc, text, bold_prefix=""):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font_style(r1, bold=True)
    r2 = p.add_run(text)
    set_font_style(r2)
    return p

def apply_quote_box(doc, text):
    """Add a callout box with a left border and shaded background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    
    # Left border and light background (using XML hacks)
    pBdr = OxmlElement('w:pBdr')
    leftBdr = OxmlElement('w:left')
    leftBdr.set(qn('w:val'), 'single')
    leftBdr.set(qn('w:sz'), '24')  # 3pt
    leftBdr.set(qn('w:space'), '12')
    leftBdr.set(qn('w:color'), '6366F1')  # Indigo left border
    pBdr.append(leftBdr)
    p._p.get_or_add_pPr().append(pBdr)
    
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FF"/>')  # Very light indigo tint
    p._p.get_or_add_pPr().append(shd)
    
    run = p.add_run(text)
    set_font_style(run, italic=True, size=10, color_rgb=RGBColor(31, 58, 86))
    return p

# ─────────────────────────────────────────────────────────────────────────────
# 1. GENERATE ROUND2_EVENING_SCRIPT.DOCX
# ─────────────────────────────────────────────────────────────────────────────
def build_round2_script():
    doc = Document()
    
    # Standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header
    p = add_paragraph_styled(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("AttriSense AI — Round 2 Evening Presentation Script")
    set_font_style(r, size=22, bold=True, color_rgb=RGBColor(31, 58, 86))
    
    p2 = add_paragraph_styled(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    r2 = p2.add_run("LTTS Global Hackathon 2026 | Intermediate Jury Round")
    set_font_style(r2, size=12, italic=True, color_rgb=RGBColor(99, 102, 241))

    apply_quote_box(doc, 
        "Context: This is your 2nd of 3 jury rounds. You have built the working system. "
        "The goal tonight is NOT a full technical demo — it is to convince the jury of the problem, your approach, and your unique differentiation. "
        "Save the live AI agents chat demo for the final round tomorrow. Focus tonight on the working dashboard, the business case, and proof of concept."
    )

    add_heading_styled(doc, "1. Team Role Assignment & Pitch Strategy", 1)
    add_paragraph_styled(doc, 
        "Our team consists of 4 members: 1 Senior CSE Engineer (5+ years experience) and 3 Graduate Engineer Trainees (GETs) "
        "from Mechanical and Electrical backgrounds. This diversity is our core pitch strength: the domain experts (GETs) "
        "identified the business problems and calibrated the model, while the CSE Engineer built the enterprise-grade AI architecture."
    )
    
    table = create_table_styled(doc, 5, 3, [1.5, 1.5, 3.5])
    populate_row(table.rows[0], ["Presenter", "Background & Role", "Speaking Focus"], bg_color="1F3A56", is_header=True)
    populate_row(table.rows[1], ["GET #1", "Mech / domain Expert", "Opening statement, problem context, LTTS & industry attrition rates, why 10% target is healthy"])
    populate_row(table.rows[2], ["GET #2", "Elec / Research Lead", "Competitive analysis, market research response, why Oracle/SAP fail, and ROI in Millions"])
    populate_row(table.rows[3], ["GET #3", "Mech / Product Lead", "HR Dashboard walk-through, live employee manual entry, and what-if scenario testing"])
    populate_row(table.rows[4], ["CSE Engineer", "CSE (5+ years exp)", "Technical stack (ADK 2.4.0, MCP, Gemini), agents architecture, security callbacks, and closing"])
    
    add_paragraph_styled(doc, "")

    add_heading_styled(doc, "2. Scripted Flow (Total 10 Minutes)", 1)
    
    add_heading_styled(doc, "Part A: Problem & Attrition Reality (GET #1 — 2.5 Min)", 2)
    add_paragraph_styled(doc, 
        "\"Good evening. I am [Name], a Graduate Engineer Trainee with LTTS. Let's start with a number: $10,000. "
        "That is the global minimum replacement cost for a single IT professional — accounting for recruiting, onboarding, training, and lost productivity. "
        "At LTTS, where attrition has historically hovered between 13% and 18% (similar to the Indian IT industry standard of 15-22%), "
        "attrition is a multi-million dollar problem. But here is the major misconception: we do NOT want an attrition rate of 0%.\""
    )
    apply_quote_box(doc, 
        "Natural Attrition is Healthy: An organization needs 8% to 12% annual attrition to support natural retirements, "
        "career transitions, performance-related exits, cost restructuring, and to invite new talent with fresh ideas. "
        "Our goal is not to stop attrition entirely. Our goal is to prevent INVOLUNTARY attrition of our top high-value performers."
    )
    
    add_heading_styled(doc, "Part B: Market Landscape & Innovation (GET #2 — 2 Min)", 2)
    add_paragraph_styled(doc, 
        "\"During our hackathon prep, we conducted market research. Big HCM players like Workday, SAP SuccessFactors, "
        "and Oracle already have flight-risk analytics. So, why build AttriSense AI?\""
    )
    apply_bullet(doc, "Explainability: Competitors show a simple risk score (e.g. 70%). We show the exact 14-parameter breakdown so HR knows exactly WHY they want to leave.", "1. ")
    apply_bullet(doc, "AI-Generated Personalized Plans: Our Retention Advisor agent designs specific intervention playbooks (actions, owners, timelines) tailored to the individual's risk drivers.", "2. ")
    apply_bullet(doc, "Real-time Entry: Instead of batch monthly runs, HR managers can input an employee on-the-fly and get instant risk calculations.", "3. ")

    add_heading_styled(doc, "Part C: Technical Architecture (CSE Engineer — 2 Min)", 2)
    add_paragraph_styled(doc, 
        "\"We built our solution using Google's new Agent Development Kit (ADK) 2.4.0. "
        "Instead of one massive, unreliable prompt, we use a Multi-Agent Orchestration pattern. "
        "Four specialist agents (Signal, Sentiment, Risk, and Retention) cooperate. "
        "Crucially, to prevent LLM hallucinations, we utilize Model Context Protocol (MCP). "
        "Our agents do not guess data — they fetch it via secure, real-time MCP tool calls. "
        "Security is baked in: our custom callback intercepts all calls to scrub PII (Aadhaar, salaries) and blocks prompt injections.\""
    )

    add_heading_styled(doc, "Part D: Live Dashboard Walkthrough (GET #3 — 2 Min)", 2)
    add_paragraph_styled(doc, 
        "\"Here is our live HR dashboard (http://127.0.0.1:18082/dashboard) currently monitoring 1,000 employees. "
        "You can see aggregate statistics, a department risk heatmap showing that Engineering is at high risk, "
        "and a searchable, paginated register of all employees. "
        "But the most powerful feature is manual entry. We can add a new employee right now. "
        "[Enter a sample: name='Test Entry', department='Engineering', tenure=3.5 years, low salary percentile, 12 job searches]. "
        "Look at that — instantly, the system calculates 75.6% HIGH Risk and highlights the compensation gap as the primary driver.\""
    )

    add_heading_styled(doc, "Part E: Closing & Q&A Prep (CSE Engineer — 1.5 Min)", 2)
    add_paragraph_styled(doc, 
        "\"In conclusion, AttriSense AI shifts HR from reactive damage control to proactive retention. "
        "Even a minor 10% reduction in LTTS's involuntary attrition translates to over $3.4 Million in annual savings. "
        "We are now ready for your questions.\""
    )

    add_heading_styled(doc, "3. High-Probability Jury Questions & Answers", 1)
    
    p_q1 = add_paragraph_styled(doc, "")
    p_q1.add_run("Q: Since Workday/SAP already do this, why would LTTS buy/build this?\n").bold = True
    p_q1.add_run(
        "A: Cost and Customization. SAP/Workday charge massive annual enterprise license fees. "
        "Additionally, those systems are closed black boxes. AttriSense AI provides transparent parameter weights "
        "and fully automated, custom retention actions. It integrates in days via the MCP layer rather than months of implementation."
    )

    p_q2 = add_paragraph_styled(doc, "")
    p_q2.add_run("Q: With 14 parameters, how do you handle data privacy?\n").bold = True
    p_q2.add_run(
        "A: Privacy is built-in. First, the data resides locally on the MCP server and never leaves the organization. "
        "Second, our before_model_callback strips all PII (names, specific salary values, identifiers) before sending the text "
        "to the Gemini API. The LLM only sees masked tokens and logical relationships."
    )

    p_q3 = add_paragraph_styled(doc, "")
    p_q3.add_run("Q: Why did you choose a Multi-Agent system over a single LLM prompt?\n").bold = True
    p_q3.add_run(
        "A: Single prompts struggle to balance complex signal evaluation, sentiment classification, math calculations, and playbook formulation. "
        "By dividing the labor among 4 specialist agents (e.g. Signal Agent handles parameters, Sentiment Agent handles NLP), "
        "each agent remains highly accurate, can be tested independently, and the overall context window stays clean, preventing model confusion."
    )

    doc.save("ROUND2_EVENING_SCRIPT.docx")
    print("Saved ROUND2_EVENING_SCRIPT.docx")

# ─────────────────────────────────────────────────────────────────────────────
# 2. GENERATE SUBMISSION_WRITEUP.DOCX
# ─────────────────────────────────────────────────────────────────────────────
def build_submission_writeup():
    doc = Document()
    
    # Margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Title
    p = add_paragraph_styled(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("AttriSense AI — Hackathon Submission Writeup")
    set_font_style(r, size=22, bold=True, color_rgb=RGBColor(31, 58, 86))
    
    p2 = add_paragraph_styled(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    r2 = p2.add_run("LTTS Global Hackathon 2026 | AI/ML Track")
    set_font_style(r2, size=12, italic=True, color_rgb=RGBColor(99, 102, 241))

    add_heading_styled(doc, "1. Project Title", 1)
    add_paragraph_styled(doc, "AttriSense AI: Predictive Employee Attrition Intelligence and Early Leaver Detection Platform")

    add_heading_styled(doc, "2. Problem Statement", 1)
    add_paragraph_styled(doc, 
        "Employee attrition costs organizations 6 to 9 months of an employee's salary per departure (SHRM benchmark). "
        "For an organization like LTTS with over 23,000 employees, an average annual attrition rate of 15% results in "
        "approximately 3,450 exits yearly. With replacement costs ranging from $10,000 to $15,000 per professional, "
        "this creates a $34.5M to $51.7M annual financial drain."
    )
    add_paragraph_styled(doc, 
        "Traditional HR approaches are reactive, occurring only after resignation. Early indicators of exit "
        "— such as increasing absenteeism, salary stagnation, lack of recognition, and declining meeting participation "
        "— typically manifest 3 to 6 months prior. AttriSense AI proactively intercepts these signals to generate early retention strategies."
    )

    add_heading_styled(doc, "3. Proposed Solution", 1)
    add_paragraph_styled(doc, 
        "AttriSense AI is an AI-powered predictive analytics platform. It leverages a multi-agent architecture to continuously "
        "evaluate 14 behavioral and operational signals, generating a risk percentage and an actionable, customized retention playbook."
    )
    
    table = create_table_styled(doc, 7, 3, [1.5, 2.5, 2.5])
    populate_row(table.rows[0], ["Dimension", "Traditional HR Tools", "AttriSense AI"], bg_color="1F3A56", is_header=True)
    populate_row(table.rows[1], ["Detection Mode", "Reactive (post-resignation)", "Proactive (3-6 months early)"])
    populate_row(table.rows[2], ["Analysis Type", "Rule-based report / Static KPIs", "Multi-Agent AI reasoning"])
    populate_row(table.rows[3], ["Key Signals Evaluated", "Basic stats (tenure, performance)", "14 behavioral & operational signals"])
    populate_row(table.rows[4], ["Intervention Approach", "Generic company-wide policies", "Personalized action plan per employee"])
    populate_row(table.rows[5], ["Security Standards", "Basic database auth", "PII Scrubbing, Prompt Injection defense, RBAC"])
    populate_row(table.rows[6], ["Data Freshness", "Monthly batch reviews", "Real-time updates + manual employee entry"])

    add_paragraph_styled(doc, "")

    add_heading_styled(doc, "4. Technical Implementation Details", 1)
    
    add_heading_styled(doc, "4.1 Multi-Agent Orchestration Flow", 2)
    add_paragraph_styled(doc, 
        "Built on Google Agent Development Kit (ADK) 2.4.0, the orchestrator routes queries through a pipeline of specialists:"
    )
    apply_bullet(doc, "Signal Detection Agent: Retrieves 14 parameters from the Model Context Protocol (MCP) server, flags warnings.", "• ")
    apply_bullet(doc, "Sentiment Analysis Agent: Performs NLP evaluations on employee survey feedback, identifying themes like burnout or manager conflict.", "• ")
    apply_bullet(doc, "Risk Scoring Agent: Runs the mathematical weighted algorithm to place the employee in Critical, High, Moderate, or Low tiers.", "• ")
    apply_bullet(doc, "Retention Advisor Agent: Automatically generates an action plan with timelines and roles.", "• ")

    add_heading_styled(doc, "4.2 The 14-Parameter Risk Model", 2)
    table_weights = create_table_styled(doc, 9, 3, [2.5, 1.0, 3.0])
    populate_row(table_weights.rows[0], ["Parameter Category", "Weight", "HR Rationale / Signal Source"], bg_color="1F3A56", is_header=True)
    populate_row(table_weights.rows[1], ["Absenteeism Rate", "12%", "Physical disengagement tracker (Monthly)"])
    populate_row(table_weights.rows[2], ["Sentiment Score & Mood Index", "15%", "Emotional disengagement from surveys (Quarterly)"])
    populate_row(table_weights.rows[3], ["Performance Trend", "10%", "Declining performance often indicates flight risk (Quarterly)"])
    populate_row(table_weights.rows[4], ["Compensation Gap", "10%", "Market pay parity comparison (Annual)"])
    populate_row(table_weights.rows[5], ["Internal Job Searches", "8%", "Portal logs indicating active internal exploration (Real-time)"])
    populate_row(table_weights.rows[6], ["Manager Scorecard", "8%", "Interpersonal dynamics, 360 review score (Quarterly)"])
    populate_row(table_weights.rows[7], ["Promotion Gap", "8%", "Career stagnation timeline (Continuous)"])
    populate_row(table_weights.rows[8], ["Collaboration & Meeting indices", "12%", "Social disengagement signals (Weekly/Monthly)"])

    add_paragraph_styled(doc, "")

    add_heading_styled(doc, "4.3 Security & PII Protection", 2)
    add_paragraph_styled(doc, 
        "Security is handled by an active callback prior to model calls. PII (Aadhaar cards, salaries) are scrubbed. "
        "Role-Based Access Control limits visibility: line managers cannot view salary percentages, while HR admins have full authorization. "
        "Prompt injection attacks are automatically detected and logged."
    )

    add_heading_styled(doc, "5. Business Impact & Scalability", 1)
    add_paragraph_styled(doc, 
        "By focusing on retaining high-value employees rather than aiming for an unrealistic 0% attrition, "
        "the application optimizes HR interventions. In a moderate scenario of preventing 10% of unplanned exits, "
        "LTTS stands to save between $3.4 Million and $5.1 Million annually. The architecture is model-agnostic "
        "and can be scaled across any enterprise using standard database connectors via the MCP server."
    )

    doc.save("SUBMISSION_WRITEUP.docx")
    print("Saved SUBMISSION_WRITEUP.docx")

# ─────────────────────────────────────────────────────────────────────────────
# 3. GENERATE FACTS_REFERENCE.DOCX
# ─────────────────────────────────────────────────────────────────────────────
def build_facts_reference():
    doc = Document()
    
    # Margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Title
    p = add_paragraph_styled(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("AttriSense AI — Key Facts & Q&A Reference Document")
    set_font_style(r, size=20, bold=True, color_rgb=RGBColor(31, 58, 86))
    
    p2 = add_paragraph_styled(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    r2 = p2.add_run("All Jury Rounds Reference Guide | LTTS Global Hackathon 2026")
    set_font_style(r2, size=11, italic=True, color_rgb=RGBColor(99, 102, 241))

    add_heading_styled(doc, "1. Industry & LTTS Attrition Benchmarks", 1)
    add_paragraph_styled(doc, 
        "India's IT and Engineering services industries face some of the highest attrition rates globally. "
        "LTTS's annual attrition historically fluctuates between 13% and 18%, aligning with peers. "
        "Industry benchmarks indicate typical annual attrition is 15-22% in IT, 12-18% in Engineering, and 8-12% in Manufacturing."
    )
    
    add_heading_styled(doc, "The Attrition Baseline Target: Why 10%?", 2)
    add_paragraph_styled(doc, 
        "A healthy company should NOT target 0% attrition. An optimal annual attrition rate is 8% to 12%. "
        "This rate accommodates natural exits (retirements), restructures, performance-based exits, and allows "
        "fresh talent to join the company. AttriSense AI focuses specifically on reducing involuntary attrition — "
        "preventing highly rated, high-performance staff from leaving due to addressable grievances."
    )

    add_heading_styled(doc, "2. Financial ROI Model (USD Millions)", 1)
    add_paragraph_styled(doc, 
        "According to the Society for Human Resource Management (SHRM), replacing an employee costs 50% to 200% of their annual salary. "
        "For LTTS (approx. 23,000 employees), a 15% attrition rate represents roughly 3,450 exits annually. "
        "Assuming a replacement cost of $10,000 to $15,000 per professional, the annual attrition expense is $34.5M to $51.7M."
    )
    
    table_roi = create_table_styled(doc, 4, 3, [2.5, 2.0, 2.0])
    populate_row(table_roi.rows[0], ["Exits Prevented Scenario", "Exits Saved / Year", "Estimated Annual Savings (USD)"], bg_color="1F3A56", is_header=True)
    populate_row(table_roi.rows[1], ["Conservative (Save 5%)", "172 employees", "$1.72M – $2.58M"])
    populate_row(table_roi.rows[2], ["Moderate (Save 10%)", "345 employees", "$3.45M – $5.17M"])
    populate_row(table_roi.rows[3], ["Optimistic (Save 15%)", "517 employees", "$5.17M – $7.75M"])

    add_paragraph_styled(doc, "")

    add_heading_styled(doc, "3. Technical Parameters: Frequencies & Thresholds", 1)
    add_paragraph_styled(doc, 
        "Frequency represents how often each metrics database is updated. "
        "Threshold represents the numerical trigger value that flags risk."
    )
    
    table_params = create_table_styled(doc, 15, 5, [2.2, 0.8, 1.0, 1.2, 1.3])
    populate_row(table_params.rows[0], ["Parameter", "Weight", "Frequency", "Safe Threshold", "Warning Threshold"], bg_color="1F3A56", is_header=True)
    populate_row(table_params.rows[1], ["Absenteeism Rate", "12%", "Monthly", "< 5% absence", "> 15% absence"])
    populate_row(table_params.rows[2], ["Sentiment Score", "10%", "Quarterly", "> +0.2 NLP score", "< -0.3 NLP score"])
    populate_row(table_params.rows[3], ["Performance Trend", "10%", "Quarterly", "High/Improving", "Declining/Low Performer"])
    populate_row(table_params.rows[4], ["Compensation Gap", "10%", "Annual", "> 60th percentile", "< 40th percentile"])
    populate_row(table_params.rows[5], ["Internal Job Search", "8%", "Real-time", "0-2 searches/mo", "> 5 searches/mo"])
    populate_row(table_params.rows[6], ["Manager Scorecard", "8%", "Quarterly", "> 7.0 / 10", "< 5.0 / 10"])
    populate_row(table_params.rows[7], ["Promotion Gap", "8%", "Continuous", "< 2 years", "> 3 years"])
    populate_row(table_params.rows[8], ["Meeting Attendance", "6%", "Weekly", "> 70%", "< 50%"])
    populate_row(table_params.rows[9], ["Collaboration Index", "6%", "Monthly", "> 70%", "< 40%"])
    populate_row(table_params.rows[10], ["Training Activity", "6%", "Monthly", "> 4 courses/yr", "0 courses in 6 mo"])
    populate_row(table_params.rows[11], ["Awards Count", "6%", "Continuous", ">= 2 awards/yr", "0 awards in 12 mo"])
    populate_row(table_params.rows[12], ["Location Mismatch", "5%", "Once/Annual", "Local", "Relocated Fresher"])
    populate_row(table_params.rows[13], ["Cross-functional Work", "3%", "Monthly", "Active", "None in 6 months"])
    populate_row(table_params.rows[14], ["Team Social Outings", "2%", "Monthly", ">= 4 / year", "0 in 6 months"])

    add_paragraph_styled(doc, "")

    add_heading_styled(doc, "4. Competitive Landscape Analysis", 1)
    add_paragraph_styled(doc, 
        "While flight-risk analytics exist in mature suites (Workday, SAP, Oracle, Visier, IBM), "
        "AttriSense AI differentiates itself through three features:"
    )
    apply_bullet(doc, "Explainability: We provide exact parameter breakdowns rather than a simple probability percentage.", "1. ")
    apply_bullet(doc, "Personalized Action Plans: Our Retention Advisor agent automatically creates playbooks mapping actions, owners, and timelines to address specific drivers.", "2. ")
    apply_bullet(doc, "Real-time Manual Entry: Allows line managers to run ad-hoc simulations and immediate assessments for team members.", "3. ")

    add_heading_styled(doc, "5. Technical and Business Q&A Cheat Sheet", 1)
    
    p_q1 = add_paragraph_styled(doc, "")
    p_q1.add_run("Q: What is the ROI in Millions of Dollars?\n").bold = True
    p_q1.add_run("A: A 10% reduction in involuntary attrition saves LTTS approximately $3.45 Million to $5.17 Million annually, against an annual operational API cost of less than $12,000.")

    p_q2 = add_paragraph_styled(doc, "")
    p_q2.add_run("Q: How does the system handle employee data security?\n").bold = True
    p_q2.add_run("A: All PII data is scrubbed locally by the before_model_callback interceptor before it reaches the model API. Salaries and private identifiers are masked or tokenized. Role-based access control prevents unauthorized manager viewing of salary figures.")

    p_q3 = add_paragraph_styled(doc, "")
    p_q3.add_run("Q: How can we connect this system to real HR databases?\n").bold = True
    p_q3.add_run("A: The application utilizes the standardized Model Context Protocol (MCP) data layer. To swap synthetic data for Workday, SAP, or a SQL server, we only update the database connector functions inside the MCP server file without touching the agent configurations.")

    doc.save("FACTS_REFERENCE.docx")
    print("Saved FACTS_REFERENCE.docx")

if __name__ == "__main__":
    build_round2_script()
    build_submission_writeup()
    build_facts_reference()
    print("All Word documents generated successfully!")
